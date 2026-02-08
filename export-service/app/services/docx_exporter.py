"""DOCX exporter for query sessions."""

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxExporter:
    """Exports a query session as a Microsoft Word (.docx) document."""

    def export_session(self, session_data: dict, workspace_name: str) -> BytesIO:
        """Generate a DOCX file from session data.

        Args:
            session_data: Dict containing session metadata and messages.
            workspace_name: The name of the workspace the session belongs to.

        Returns:
            BytesIO buffer containing the generated DOCX file.
        """
        doc = Document()
        session_title = session_data.get("title", "Untitled Session")
        messages = session_data.get("messages", [])
        export_date = datetime.now(timezone.utc).strftime("%B %d, %Y at %H:%M UTC")

        # ── Title ──────────────────────────────────────────────────────
        title_para = doc.add_heading(level=0)
        title_run = title_para.add_run(f"{workspace_name} - {session_title}")
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # ── Subtitle / date ───────────────────────────────────────────
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(f"Exported on {export_date}")
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()  # spacer

        # ── Message pairs ─────────────────────────────────────────────
        pairs = self._extract_message_pairs(messages)
        for question, answer, citations in pairs:
            # Question heading
            q_heading = doc.add_heading(level=2)
            q_run = q_heading.add_run(f"Q: {question}")
            q_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # blue
            q_run.font.size = Pt(14)

            # Answer body
            if answer:
                for paragraph_text in answer.split("\n\n"):
                    stripped = paragraph_text.strip()
                    if stripped:
                        doc.add_paragraph(stripped, style="Normal")

            # Citations / sources
            if citations:
                sources_para = doc.add_paragraph()
                sources_run = sources_para.add_run("Sources:")
                sources_run.bold = True
                sources_run.font.size = Pt(10)

                for idx, cite in enumerate(citations, start=1):
                    doc_name = cite.get("document_name", "Unknown Document")
                    page = cite.get("page", "N/A")
                    section = cite.get("section", "")

                    cite_text = f"{idx}. {doc_name}, Page {page}"
                    if section:
                        cite_text += f", Section: {section}"

                    cite_para = doc.add_paragraph(cite_text, style="Normal")
                    for run in cite_para.runs:
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Horizontal rule between pairs
            doc.add_paragraph("_" * 60)

        # ── Serialize to bytes ────────────────────────────────────────
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _extract_message_pairs(
        messages: list[dict],
    ) -> list[tuple[str, str, list[dict]]]:
        """Pair up user questions with assistant answers.

        Returns a list of (question, answer, citations) tuples.
        """
        pairs: list[tuple[str, str, list[dict]]] = []
        i = 0
        while i < len(messages):
            msg = messages[i]
            if msg.get("role") == "user":
                question = msg.get("content", "")
                answer = ""
                citations: list[dict] = []

                # Look for the following assistant message
                if i + 1 < len(messages) and messages[i + 1].get("role") == "assistant":
                    assistant_msg = messages[i + 1]
                    answer = assistant_msg.get("content", "")
                    citations = assistant_msg.get("citations", [])
                    i += 2
                else:
                    i += 1

                pairs.append((question, answer, citations))
            else:
                i += 1

        return pairs
